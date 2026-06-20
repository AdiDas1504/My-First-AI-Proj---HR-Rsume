from datetime import datetime
from pathlib import Path

from docx import Document


def format_list(title, items):
    """
    Format a list of items for the text report.
    """
    lines = [
        title,
        "-" * 50,
    ]

    if not items:
        lines.append("No items found.")
    else:
        for item in items:
            lines.append(f"- {item}")

    lines.append("")
    return "\n".join(lines)


def build_text_report(report, tailoring_plan, resume_text, job_text):
    """
    Build a full text report from the fit report and tailoring plan.
    """
    lines = []

    lines.append("JobFit AI Resume Tailor Report")
    lines.append("=" * 70)
    lines.append("")

    lines.append("Candidate Fit Summary")
    lines.append("-" * 70)
    lines.append(f"Fit score: {report['fit_score']}%")
    lines.append(f"Fit level: {report['fit_level']}")
    lines.append("")

    lines.append("Recommendation")
    lines.append("-" * 70)
    lines.append(report["recommendation"])
    lines.append("")

    lines.append(format_list("Matched Keywords", report["matched_keywords"]))
    lines.append(format_list("Missing or Weak Keywords", report["missing_keywords"]))

    lines.append("Resume Tailoring Plan")
    lines.append("=" * 70)
    lines.append("")

    lines.append("Application Priority")
    lines.append("-" * 70)
    lines.append(tailoring_plan["priority"])
    lines.append("")

    lines.append("Professional Summary Suggestion")
    lines.append("-" * 70)
    lines.append(tailoring_plan["summary_suggestion"])
    lines.append("")

    lines.append(format_list(
        "Missing Keywords to Consider Only If True",
        tailoring_plan["missing_keyword_suggestions"]
    ))

    lines.append(format_list(
        "Bullet Point Rewrite Guidelines",
        tailoring_plan["bullet_guidelines"]
    ))

    lines.append("Extracted Resume Text Preview")
    lines.append("=" * 70)
    lines.append(resume_text[:1500])
    lines.append("")

    lines.append("Extracted Job Requirements Preview")
    lines.append("=" * 70)
    lines.append(job_text[:1500])
    lines.append("")

    lines.append("Important Note")
    lines.append("=" * 70)
    lines.append(
        "This report should not be used to invent false experience, skills, "
        "education, achievements, or tools. Resume improvements should only "
        "reflect the candidate's real background."
    )

    return "\n".join(lines)


def save_text_report(report, tailoring_plan, resume_text, job_text):
    """
    Save the full analysis report as a TXT file.
    """
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"fit_report_{timestamp}.txt"

    report_text = build_text_report(
        report=report,
        tailoring_plan=tailoring_plan,
        resume_text=resume_text,
        job_text=job_text,
    )

    output_path.write_text(report_text, encoding="utf-8")

    return str(output_path)


def add_bullet_list(document, items):
    """
    Add a bullet list to a Word document.
    """
    if not items:
        document.add_paragraph("No items found.")
        return

    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def save_word_report(report, tailoring_plan, resume_text, job_text):
    """
    Save the full analysis report as a Word document.
    """
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = output_dir / f"fit_report_{timestamp}.docx"

    document = Document()

    document.add_heading("JobFit AI Resume Tailor Report", level=1)

    document.add_heading("Candidate Fit Summary", level=2)
    document.add_paragraph(f"Fit score: {report['fit_score']}%")
    document.add_paragraph(f"Fit level: {report['fit_level']}")

    document.add_heading("Recommendation", level=2)
    document.add_paragraph(report["recommendation"])

    document.add_heading("Matched Keywords", level=2)
    add_bullet_list(document, report["matched_keywords"])

    document.add_heading("Missing or Weak Keywords", level=2)
    add_bullet_list(document, report["missing_keywords"])

    document.add_heading("Resume Tailoring Plan", level=1)

    document.add_heading("Application Priority", level=2)
    document.add_paragraph(tailoring_plan["priority"])

    document.add_heading("Professional Summary Suggestion", level=2)
    document.add_paragraph(tailoring_plan["summary_suggestion"])

    document.add_heading("Missing Keywords to Consider Only If True", level=2)
    add_bullet_list(document, tailoring_plan["missing_keyword_suggestions"])

    document.add_heading("Bullet Point Rewrite Guidelines", level=2)
    add_bullet_list(document, tailoring_plan["bullet_guidelines"])

    document.add_heading("Extracted Resume Text Preview", level=2)
    document.add_paragraph(resume_text[:1500])

    document.add_heading("Extracted Job Requirements Preview", level=2)
    document.add_paragraph(job_text[:1500])

    document.add_heading("Important Note", level=2)
    document.add_paragraph(
        "This report should not be used to invent false experience, skills, "
        "education, achievements, or tools. Resume improvements should only "
        "reflect the candidate's real background."
    )

    document.save(output_path)

    return str(output_path)